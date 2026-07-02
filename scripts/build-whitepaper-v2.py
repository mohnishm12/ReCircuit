from __future__ import annotations

import json
import shutil
import textwrap
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Flowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "src" / "data" / "research-whitepaper-v2.json"
PUBLIC_DIR = ROOT / "public" / "research"
DIAGRAM_DIR = PUBLIC_DIR / "diagrams"
QA_DIR = ROOT / "tmp" / "research-build"

INK = "#F5F7F2"
MUTED = "#9CA3AF"
GREEN = "#14F195"
CYAN = "#7DF9FF"
BLUE = "#6EA8FE"
AMBER = "#F8C14A"
SURFACE = "#0B0F0E"
PANEL = "#111816"

DIAGRAM_NODES = {
    "system-architecture": ["Frontend", "API Layer", "Backend Services", "Data Pipeline", "AI Layer", "Evidence Store", "Reporting Exports"],
    "data-flow": ["Intake", "Validation", "Model Review", "Decision", "Human Override Log", "ESG Evidence Packet"],
    "circular-economy": ["Electronics", "Collection", "Triage", "Reuse + Repair", "Component Recovery", "Recycling", "New Materials"],
    "ai-pipeline": ["Images", "Metadata", "Computer Vision", "Forecasting", "Recommendation", "Carbon + ESG Methods", "Confidence Review"],
    "enterprise-workflow": ["Consumers", "Businesses", "Recyclers", "Asset Intelligence Graph", "ESG Reports", "Government Views"],
}


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def clean(value: str) -> str:
    return " ".join(str(value).replace("\u2013", "-").replace("\u2014", "-").split())


def wrap(value: str, width: int = 96) -> str:
    return "\n".join(textwrap.wrap(clean(value), width=width))


def ensure_dirs() -> None:
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def ieee_reference(ref: dict) -> str:
    return f"[{ref['id']}] {ref['label']}, \"{ref['title']},\" {ref['year']}. [Online]. Available: {ref['url']}"


def build_markdown(data: dict) -> str:
    meta = data["meta"]
    lines: list[str] = [
        f"# {meta['title']}",
        "",
        f"## {meta['subtitle']}",
        "",
        f"**Author:** {meta['author']}  ",
        f"**Role:** {meta['authorRole']}  ",
        f"**Version:** {meta['version']}  ",
        f"**Publication Date:** {meta['publicationDate']}  ",
        f"**Status:** {meta['status']}",
        "",
        "## Abstract",
        "",
    ]
    for paragraph in data["abstract"]:
        lines += [wrap(paragraph), ""]

    lines += ["## Executive Summary", ""]
    for item in data["executiveSummary"]:
        lines += [f"### {item['question']}", "", wrap(item["answer"]), ""]

    lines += ["## Original Diagrams", ""]
    for diagram in data["diagrams"]:
        lines += [
            f"### {diagram['title']}",
            "",
            wrap(diagram["description"]),
            "",
            f"![{diagram['title']}](diagrams/{diagram['id']}.svg)",
            "",
        ]

    for section in data["sections"]:
        lines += [f"## {section['title']}", "", f"_{section['kicker']}_", ""]
        for paragraph in section["paragraphs"]:
            lines += [wrap(paragraph), ""]
        if section.get("currentStatus"):
            lines += ["**Current status**", ""]
            for entry in section["currentStatus"]:
                lines += [f"- {wrap(entry)}"]
            lines += [""]
        if section.get("futureCapabilities"):
            lines += ["**Proposed future capabilities**", ""]
            for entry in section["futureCapabilities"]:
                lines += [f"- {wrap(entry)}"]
            lines += [""]

    lines += ["## System Architecture", ""]
    for subsystem in data["architecture"]:
        lines += [
            f"### {subsystem['name']}",
            "",
            f"**Purpose:** {wrap(subsystem['purpose'])}",
            "",
            f"**Inputs:** {wrap(subsystem['inputs'])}",
            "",
            f"**Outputs:** {wrap(subsystem['outputs'])}",
            "",
            f"**Current status:** {wrap(subsystem['currentStatus'])}",
            "",
            f"**Future work:** {wrap(subsystem['futureWork'])}",
            "",
        ]

    lines += ["## Enterprise Workflows", ""]
    for workflow in data["workflows"]:
        lines += [
            f"### {workflow['stakeholder']}",
            "",
            wrap(workflow["workflow"]),
            "",
            f"**Value:** {wrap(workflow['value'])}",
            "",
        ]

    lines += ["## Roadmap", ""]
    for item in data["roadmap"]:
        lines += [f"- **{item['period']} - {item['stage']}:** {wrap(item['summary'])}"]
    lines += [""]

    lines += ["## Risk Register", ""]
    for risk in data["risks"]:
        lines += [f"- **{risk['category']}:** {wrap(risk['risk'])} Mitigation: {wrap(risk['mitigation'])}"]
    lines += [""]

    lines += ["## Appendices", "", "### Glossary", ""]
    for item in data["appendices"]["glossary"]:
        lines += [f"- **{item['term']}:** {wrap(item['definition'])}"]
    lines += ["", "### Future Research", ""]
    for item in data["appendices"]["futureResearch"]:
        lines += [f"- {wrap(item)}"]

    lines += ["", "## References", ""]
    for ref in data["references"]:
        lines += [ieee_reference(ref), ""]
    return "\n".join(lines)


def write_companion_docs(data: dict) -> None:
    meta = data["meta"]
    markdown = build_markdown(data)
    write_text(PUBLIC_DIR / "recircuit-research-whitepaper-v2.0.md", markdown)

    summary_lines = [f"# Executive Summary - {meta['title']}", ""]
    for item in data["executiveSummary"]:
        summary_lines += [f"## {item['question']}", "", wrap(item["answer"]), ""]
    write_text(PUBLIC_DIR / "executive-summary.md", "\n".join(summary_lines))

    architecture_lines = [
        "# Technical Architecture Companion",
        "",
        "This companion describes the ReCircuit platform architecture as a research and prototype direction. It does not claim production deployment.",
        "",
    ]
    for subsystem in data["architecture"]:
        architecture_lines += [
            f"## {subsystem['name']}",
            "",
            f"Purpose: {wrap(subsystem['purpose'])}",
            "",
            f"Inputs: {wrap(subsystem['inputs'])}",
            "",
            f"Outputs: {wrap(subsystem['outputs'])}",
            "",
            f"Current status: {wrap(subsystem['currentStatus'])}",
            "",
            f"Future work: {wrap(subsystem['futureWork'])}",
            "",
        ]
    write_text(PUBLIC_DIR / "technical-architecture-companion.md", "\n".join(architecture_lines))

    changelog = [
        "# Change Log",
        "",
        "## Version 2.0 - 2026",
        "",
        "- Rewrote the short founder brief into a structured research whitepaper.",
        "- Added global context, system architecture, AI methods, enterprise workflows, roadmap, risks, founder biography, references, and appendices.",
        "- Added explicit current-status vs proposed-capability distinctions.",
        "- Removed any implication of funding, revenue, customers, investors, employees, partnerships, awards, deployments, or performance metrics.",
        "- Added original SVG diagrams, PDF, DOCX, Markdown source, citation metadata, and research portal pages.",
        "",
        "## Version 0.1 - 2026",
        "",
        "- Initial concise founder research note.",
    ]
    write_text(PUBLIC_DIR / "change-log.md", "\n".join(changelog))

    write_text(PUBLIC_DIR / "version-history.json", json.dumps(data["versionHistory"], indent=2))

    citation_cff = f"""
cff-version: 1.2.0
message: "If you use this publication, cite it using the metadata below."
title: "{meta['title']}: {meta['subtitle']}"
version: "{meta['version']}"
date-released: "2026"
authors:
  - family-names: "M"
    given-names: "Mohnish"
url: "{meta['canonicalUrl']}"
abstract: "{clean(data['abstract'][0])}"
license: "All rights reserved"
"""
    write_text(PUBLIC_DIR / "citation.cff", citation_cff)

    citation_json = {
        "@context": "https://schema.org",
        "@type": "ScholarlyArticle",
        "name": f"{meta['title']}: {meta['subtitle']}",
        "author": {"@type": "Person", "name": meta["author"]},
        "version": meta["version"],
        "datePublished": meta["publicationDate"],
        "url": meta["canonicalUrl"],
        "abstract": " ".join(data["abstract"]),
        "publisher": {"@type": "Organization", "name": "ReCircuit Research"},
    }
    write_text(PUBLIC_DIR / "citation.json", json.dumps(citation_json, indent=2))


def svg_rect(x: int, y: int, w: int, h: int, label: str, stroke: str = GREEN, fill: str = PANEL) -> str:
    safe = clean(label)
    return f"""
  <rect x="{x}" y="{y}" width="{w}" height="{h}" rx="10" fill="{fill}" stroke="{stroke}" stroke-width="1.4"/>
  <text x="{x + w / 2}" y="{y + h / 2 - 4}" text-anchor="middle" fill="{INK}" font-size="16" font-family="Inter, Arial" font-weight="700">{safe}</text>
"""


def svg_arrow(x1: int, y1: int, x2: int, y2: int, stroke: str = CYAN) -> str:
    return f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{stroke}" stroke-width="1.4" marker-end="url(#arrow)" opacity="0.85"/>'


def write_svg(path: Path, title: str, body: str) -> None:
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1200 680" role="img" aria-labelledby="title desc">
  <title id="title">{title}</title>
  <desc id="desc">Original ReCircuit research diagram.</desc>
  <defs>
    <marker id="arrow" viewBox="0 0 10 10" refX="8" refY="5" markerWidth="6" markerHeight="6" orient="auto-start-reverse">
      <path d="M 0 0 L 10 5 L 0 10 z" fill="{CYAN}"/>
    </marker>
    <linearGradient id="bg" x1="0" x2="1">
      <stop offset="0" stop-color="#050706"/>
      <stop offset="1" stop-color="#101613"/>
    </linearGradient>
  </defs>
  <rect width="1200" height="680" fill="url(#bg)"/>
  <text x="60" y="72" fill="{INK}" font-size="30" font-family="Inter, Arial" font-weight="800">{title}</text>
  <text x="60" y="105" fill="{MUTED}" font-size="15" font-family="Inter, Arial">ReCircuit Research Whitepaper v2.0</text>
{body}
</svg>
"""
    write_text(path, svg)


def build_svgs() -> None:
    write_svg(
        DIAGRAM_DIR / "system-architecture.svg",
        "System Architecture",
        "\n".join(
            [
                svg_rect(60, 170, 190, 74, "Frontend", BLUE),
                svg_arrow(250, 207, 335, 207),
                svg_rect(335, 170, 190, 74, "API Layer", CYAN),
                svg_arrow(525, 207, 610, 207),
                svg_rect(610, 170, 230, 74, "Backend Services", GREEN),
                svg_arrow(725, 244, 725, 318),
                svg_rect(610, 318, 230, 74, "Data Pipeline", AMBER),
                svg_arrow(840, 355, 930, 355),
                svg_rect(930, 318, 205, 74, "Evidence Store", BLUE),
                svg_arrow(725, 392, 725, 465),
                svg_rect(610, 465, 230, 74, "Reporting Exports", CYAN),
                svg_rect(930, 170, 205, 74, "AI Layer", GREEN),
                svg_arrow(930, 207, 840, 207),
                '<text x="60" y="590" fill="#9CA3AF" font-size="16" font-family="Inter, Arial">Current: research + prototype architecture. Future: secure services, model evaluation, integrations, audit exports.</text>',
            ]
        ),
    )
    write_svg(
        DIAGRAM_DIR / "data-flow.svg",
        "Data Flow",
        "\n".join(
            [
                svg_rect(72, 250, 160, 72, "Intake", GREEN),
                svg_arrow(232, 286, 310, 286),
                svg_rect(310, 250, 170, 72, "Validation", CYAN),
                svg_arrow(480, 286, 558, 286),
                svg_rect(558, 250, 190, 72, "Model Review", BLUE),
                svg_arrow(748, 286, 826, 286),
                svg_rect(826, 250, 180, 72, "Decision", AMBER),
                svg_arrow(916, 322, 916, 410),
                svg_rect(790, 410, 252, 72, "ESG Evidence Packet", GREEN),
                svg_arrow(558, 286, 558, 410),
                svg_rect(445, 410, 226, 72, "Human Override Log", CYAN),
                '<text x="72" y="560" fill="#9CA3AF" font-size="16" font-family="Inter, Arial">Evidence never disappears: raw inputs, model outputs, decisions, and overrides remain traceable.</text>',
            ]
        ),
    )
    write_svg(
        DIAGRAM_DIR / "circular-economy.svg",
        "Circular Economy Loop",
        "\n".join(
            [
                '<circle cx="600" cy="350" r="210" fill="none" stroke="#20342D" stroke-width="22"/>',
                svg_rect(500, 116, 200, 64, "Electronics", GREEN),
                svg_rect(785, 245, 190, 64, "Collection", CYAN),
                svg_rect(735, 485, 210, 64, "Reuse + Repair", BLUE),
                svg_rect(455, 552, 250, 64, "Component Recovery", AMBER),
                svg_rect(210, 415, 190, 64, "Recycling", GREEN),
                svg_rect(225, 210, 200, 64, "New Materials", CYAN),
                '<path d="M690 150 C900 205 965 415 845 520" fill="none" stroke="#7DF9FF" stroke-width="3" marker-end="url(#arrow)"/>',
                '<path d="M700 585 C455 645 220 520 260 300" fill="none" stroke="#14F195" stroke-width="3" marker-end="url(#arrow)"/>',
                '<text x="600" y="346" text-anchor="middle" fill="#F5F7F2" font-size="28" font-family="Inter, Arial" font-weight="800">Material</text>',
                '<text x="600" y="380" text-anchor="middle" fill="#9CA3AF" font-size="18" font-family="Inter, Arial">intelligence layer</text>',
            ]
        ),
    )
    write_svg(
        DIAGRAM_DIR / "ai-pipeline.svg",
        "AI Pipeline",
        "\n".join(
            [
                svg_rect(70, 210, 160, 68, "Images", GREEN),
                svg_rect(70, 310, 160, 68, "Metadata", BLUE),
                svg_arrow(230, 244, 330, 278),
                svg_arrow(230, 344, 330, 312),
                svg_rect(330, 250, 170, 82, "Computer Vision", CYAN),
                svg_arrow(500, 291, 575, 291),
                svg_rect(575, 250, 170, 82, "Forecasting", AMBER),
                svg_arrow(745, 291, 820, 291),
                svg_rect(820, 250, 210, 82, "Recommendation", GREEN),
                svg_arrow(925, 332, 925, 415),
                svg_rect(820, 415, 210, 82, "Confidence Review", CYAN),
                svg_rect(330, 415, 250, 82, "Carbon + ESG Methods", BLUE),
                svg_arrow(580, 456, 820, 456),
                '<text x="70" y="575" fill="#9CA3AF" font-size="16" font-family="Inter, Arial">AI supports classification, estimation, and explanation. It does not replace evidence or human review.</text>',
            ]
        ),
    )
    write_svg(
        DIAGRAM_DIR / "enterprise-workflow.svg",
        "Enterprise Workflow",
        "\n".join(
            [
                svg_rect(80, 170, 190, 70, "Consumers", GREEN),
                svg_rect(80, 300, 190, 70, "Businesses", BLUE),
                svg_rect(80, 430, 190, 70, "Recyclers", CYAN),
                svg_arrow(270, 205, 480, 300),
                svg_arrow(270, 335, 480, 335),
                svg_arrow(270, 465, 480, 370),
                svg_rect(480, 270, 265, 120, "Asset Intelligence Graph", AMBER),
                svg_arrow(745, 330, 915, 235),
                svg_arrow(745, 330, 915, 425),
                svg_rect(915, 200, 210, 70, "ESG Reports", GREEN),
                svg_rect(915, 390, 210, 70, "Government Views", BLUE),
                '<text x="480" y="585" fill="#9CA3AF" font-size="16" font-family="Inter, Arial">Different users, one evidence model: identity, condition, custody, pathway, carbon, and compliance context.</text>',
            ]
        ),
    )


class Rule(Flowable):
    def __init__(self, width: float, color_hex: str = GREEN, thickness: float = 1.2):
        super().__init__()
        self.width = width
        self.color_hex = color_hex
        self.thickness = thickness
        self.height = 8

    def draw(self) -> None:
        self.canv.setStrokeColor(colors.HexColor(self.color_hex))
        self.canv.setLineWidth(self.thickness)
        self.canv.line(0, 4, self.width, 4)


def pdf_styles() -> dict:
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ReTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=30,
            leading=34,
            textColor=colors.HexColor("#111111"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "subtitle": ParagraphStyle(
            "ReSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=15,
            leading=20,
            textColor=colors.HexColor("#36403A"),
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "ReH1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=22,
            textColor=colors.HexColor("#0B5E46"),
            spaceBefore=14,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "ReH2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12.5,
            leading=16,
            textColor=colors.HexColor("#24445F"),
            spaceBefore=8,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "ReBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.7,
            leading=14.1,
            alignment=TA_JUSTIFY,
            textColor=colors.HexColor("#111111"),
            spaceAfter=7,
        ),
        "lead": ParagraphStyle(
            "ReLead",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=16,
            textColor=colors.HexColor("#24312D"),
            spaceAfter=10,
        ),
        "small": ParagraphStyle(
            "ReSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10.5,
            textColor=colors.HexColor("#4A4A4A"),
            spaceAfter=4,
        ),
        "cell": ParagraphStyle(
            "ReCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=7.5,
            leading=9.5,
            textColor=colors.HexColor("#111111"),
        ),
        "cell_bold": ParagraphStyle(
            "ReCellBold",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=7.6,
            leading=9.6,
            textColor=colors.HexColor("#111111"),
        ),
    }


def pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#6F7672"))
    canvas.drawString(doc.leftMargin, 0.45 * inch, "ReCircuit Research Whitepaper v2.0")
    canvas.drawRightString(LETTER[0] - doc.rightMargin, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def bullet_list(items: Iterable[str], style: ParagraphStyle) -> ListFlowable:
    return ListFlowable(
        [ListItem(Paragraph(clean(item), style), leftIndent=14) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=7,
    )


def pdf_table(rows: list[list[str]], widths: list[float], style_map: dict) -> Table:
    table = Table(
        [[Paragraph(clean(cell), style_map["cell_bold"] if i == 0 else style_map["cell"]) for cell in row] for i, row in enumerate(rows)],
        colWidths=widths,
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F3EE")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111111")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D2CD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_diagram(title: str, nodes: list[str], styles: dict) -> KeepTogether:
    cells = [[Paragraph(node, styles["cell_bold"]) for node in nodes[:4]]]
    if len(nodes) > 4:
        cells.append([Paragraph(node, styles["cell_bold"]) for node in nodes[4:8]])
    col_count = max(len(row) for row in cells)
    for row in cells:
        while len(row) < col_count:
            row.append(Paragraph("", styles["cell"]))
    table = Table(cells, colWidths=[6.5 * inch / col_count] * col_count, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#AFC8BA")),
                ("INNERGRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#C8D2CD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    return KeepTogether([Paragraph(title, styles["h2"]), table, Spacer(1, 8)])


def build_pdf(data: dict) -> None:
    meta = data["meta"]
    path = PUBLIC_DIR / "recircuit-research-whitepaper-v2.0.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.82 * inch,
        rightMargin=0.82 * inch,
        topMargin=0.78 * inch,
        bottomMargin=0.72 * inch,
        title=f"{meta['title']}: {meta['subtitle']}",
        author=meta["author"],
    )
    styles = pdf_styles()
    story = []

    story += [Spacer(1, 1.25 * inch)]
    story += [Paragraph(meta["title"], styles["title"])]
    story += [Paragraph(meta["subtitle"], styles["subtitle"])]
    story += [Spacer(1, 0.2 * inch), Rule(5.1 * inch, "#14F195"), Spacer(1, 0.35 * inch)]
    story += [
        Paragraph(f"{meta['author']}<br/>{meta['authorRole']}<br/>Version {meta['version']} - {meta['publicationDate']}", styles["subtitle"]),
        Spacer(1, 0.5 * inch),
        Paragraph(clean(meta["status"]), styles["lead"]),
        PageBreak(),
    ]

    story += [Paragraph("Executive Summary", styles["h1"])]
    for item in data["executiveSummary"]:
        story += [Paragraph(item["question"], styles["h2"]), Paragraph(item["answer"], styles["body"])]

    story += [PageBreak(), Paragraph("Abstract", styles["h1"])]
    for paragraph in data["abstract"]:
        story += [Paragraph(paragraph, styles["lead"])]

    story += [PageBreak(), Paragraph("Original Diagrams", styles["h1"])]
    story += [
        Paragraph(
            "The website package includes original SVG diagrams. The PDF below reproduces each diagram as publication-safe vector tables.",
            styles["lead"],
        )
    ]
    for diagram in data["diagrams"]:
        story += [pdf_diagram(diagram["title"], DIAGRAM_NODES[diagram["id"]], styles)]

    for section in data["sections"]:
        story += [Paragraph(section["title"], styles["h1"]), Paragraph(section["kicker"], styles["lead"])]
        for paragraph in section["paragraphs"]:
            story += [Paragraph(paragraph, styles["body"])]
        if section.get("currentStatus"):
            story += [Paragraph("Current Status", styles["h2"]), bullet_list(section["currentStatus"], styles["body"])]
        if section.get("futureCapabilities"):
            story += [Paragraph("Proposed Future Capabilities", styles["h2"]), bullet_list(section["futureCapabilities"], styles["body"])]

    story += [PageBreak(), Paragraph("System Architecture", styles["h1"])]
    story += [
        Paragraph(
            "Each subsystem is described by purpose, inputs, outputs, current status, and future work so capability claims remain bounded.",
            styles["lead"],
        )
    ]
    arch_rows = [["Subsystem", "Purpose", "Current status", "Future work"]]
    for subsystem in data["architecture"]:
        arch_rows.append([subsystem["name"], subsystem["purpose"], subsystem["currentStatus"], subsystem["futureWork"]])
    story += [pdf_table(arch_rows, [1.05 * inch, 1.9 * inch, 1.75 * inch, 1.8 * inch], styles)]

    story += [PageBreak(), Paragraph("Enterprise Workflows", styles["h1"])]
    rows = [["Stakeholder", "Workflow", "Value"]]
    for workflow in data["workflows"]:
        rows.append([workflow["stakeholder"], workflow["workflow"], workflow["value"]])
    story += [pdf_table(rows, [1.25 * inch, 3.0 * inch, 2.25 * inch], styles)]

    story += [Paragraph("Roadmap", styles["h1"])]
    rows = [["Period", "Stage", "Milestone"]]
    for item in data["roadmap"]:
        rows.append([item["period"], item["stage"], item["summary"]])
    story += [pdf_table(rows, [1.05 * inch, 1.25 * inch, 4.2 * inch], styles)]

    story += [PageBreak(), Paragraph("Risk Register", styles["h1"])]
    rows = [["Category", "Risk", "Mitigation"]]
    for item in data["risks"]:
        rows.append([item["category"], item["risk"], item["mitigation"]])
    story += [pdf_table(rows, [1.15 * inch, 2.65 * inch, 2.7 * inch], styles)]

    story += [Paragraph("Appendix A: Glossary", styles["h1"])]
    for item in data["appendices"]["glossary"]:
        story += [Paragraph(item["term"], styles["h2"]), Paragraph(item["definition"], styles["body"])]

    story += [Paragraph("Appendix B: Future Research", styles["h1"])]
    story += [bullet_list(data["appendices"]["futureResearch"], styles["body"])]

    story += [PageBreak(), Paragraph("References", styles["h1"])]
    for ref in data["references"]:
        story += [Paragraph(ieee_reference(ref), styles["small"])]

    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width)
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(row.cells[index])
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_docx_paragraph(doc, text: str, style="Normal", bold=False, italic=False) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p)
    run = p.add_run(clean(text))
    set_run_font(run, bold=bold, italic=italic)


def add_docx_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr_cells[idx], "F4F6F9")
        paragraph = hdr_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=9, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            run = paragraph.add_run(clean(value))
            set_run_font(run, size=8.2)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_docx_diagram(doc, title: str, nodes: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(title)
    cols = min(3, len(nodes))
    rows = (len(nodes) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(nodes):
                set_cell_shading(cell, "E8EEF5" if idx % 2 == 0 else "F4F6F9")
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(nodes[idx])
                set_run_font(run, size=9.5, bold=True, color="0B2545")
            idx += 1
    set_table_geometry(table, [2.16] * cols)
    doc.add_paragraph()


def strip_docx_metadata(path: Path) -> None:
    core = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>ReCircuit Research Whitepaper</dc:title>
  <dc:subject>Building the Intelligence Layer for the Circular Economy</dc:subject>
  <dc:creator>ReCircuit Research</dc:creator>
  <cp:keywords>ReCircuit; circular economy; e-waste; AI; ESG</cp:keywords>
  <dc:description>Publication-ready ReCircuit Research Whitepaper v2.0.</dc:description>
  <cp:lastModifiedBy>ReCircuit Research</cp:lastModifiedBy>
</cp:coreProperties>"""
    temp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename == "docProps/core.xml":
                target.writestr(item, core)
            else:
                target.writestr(item, source.read(item.filename))
    temp_path.replace(path)


def build_docx(data: dict) -> None:
    meta = data["meta"]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "ReCircuit Research"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=8.5, color="666666")
    footer = section.footer.paragraphs[0]
    footer.text = "ReCircuit Research Whitepaper v2.0"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.runs[0], size=8.5, color="666666")

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cover, before=132, after=18, line=1)
    run = cover.add_run("RESEARCH WHITEPAPER")
    set_run_font(run, size=10.5, color="7A5A00", bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=8, line=1)
    run = title.add_run(meta["title"])
    set_run_font(run, size=30, color="203748", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=28, line=1.15)
    run = subtitle.add_run(meta["subtitle"])
    set_run_font(run, size=15, color="2B5163")

    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(byline, after=88, line=1.15)
    run = byline.add_run(f"{meta['author']} | {meta['authorRole']}")
    set_run_font(run, size=10.5, color="7A5A00", bold=True)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(version, after=4, line=1.15)
    run = version.add_run(f"Version {meta['version']} | Publication Date {meta['publicationDate']}")
    set_run_font(run, size=12, color="203748", bold=True)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(status, after=22, line=1.15)
    run = status.add_run(meta["status"])
    set_run_font(run, size=9.5, color="505050", italic=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_docx_paragraph(doc, "Executive Summary", style="Heading 1")
    for item in data["executiveSummary"]:
        add_docx_paragraph(doc, item["question"], style="Heading 2")
        add_docx_paragraph(doc, item["answer"])

    add_docx_paragraph(doc, "Abstract", style="Heading 1")
    for paragraph in data["abstract"]:
        add_docx_paragraph(doc, paragraph)

    for section_data in data["sections"]:
        add_docx_paragraph(doc, section_data["title"], style="Heading 1")
        add_docx_paragraph(doc, section_data["kicker"], italic=True)
        for paragraph in section_data["paragraphs"]:
            add_docx_paragraph(doc, paragraph)
        if section_data.get("currentStatus"):
            add_docx_paragraph(doc, "Current Status", style="Heading 2")
            for item in section_data["currentStatus"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(clean(item))
        if section_data.get("futureCapabilities"):
            add_docx_paragraph(doc, "Proposed Future Capabilities", style="Heading 2")
            for item in section_data["futureCapabilities"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(clean(item))

    add_docx_paragraph(doc, "Original Diagrams", style="Heading 1")
    for diagram in data["diagrams"]:
        add_docx_diagram(doc, f"Original Diagram: {diagram['title']}", DIAGRAM_NODES[diagram["id"]])

    add_docx_paragraph(doc, "System Architecture", style="Heading 1")
    add_docx_table(
        doc,
        ["Subsystem", "Purpose", "Current status", "Future work"],
        [[item["name"], item["purpose"], item["currentStatus"], item["futureWork"]] for item in data["architecture"]],
        [1.15, 1.9, 1.65, 1.8],
    )

    add_docx_paragraph(doc, "Enterprise Workflows", style="Heading 1")
    add_docx_diagram(doc, "Original Diagram: Enterprise Workflow", ["Consumers", "Businesses", "Recyclers", "Asset Intelligence Graph", "ESG Reports", "Government Views"])
    add_docx_table(
        doc,
        ["Stakeholder", "Workflow", "Value"],
        [[item["stakeholder"], item["workflow"], item["value"]] for item in data["workflows"]],
        [1.25, 3.0, 2.25],
    )

    add_docx_paragraph(doc, "Roadmap", style="Heading 1")
    add_docx_table(
        doc,
        ["Period", "Stage", "Milestone"],
        [[item["period"], item["stage"], item["summary"]] for item in data["roadmap"]],
        [1.05, 1.25, 4.2],
    )

    add_docx_paragraph(doc, "Risk Register", style="Heading 1")
    add_docx_table(
        doc,
        ["Category", "Risk", "Mitigation"],
        [[item["category"], item["risk"], item["mitigation"]] for item in data["risks"]],
        [1.15, 2.65, 2.7],
    )

    add_docx_paragraph(doc, "Appendix A: Glossary", style="Heading 1")
    for item in data["appendices"]["glossary"]:
        add_docx_paragraph(doc, item["term"], style="Heading 2")
        add_docx_paragraph(doc, item["definition"])

    add_docx_paragraph(doc, "Appendix B: Future Research", style="Heading 1")
    for item in data["appendices"]["futureResearch"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(clean(item))

    add_docx_paragraph(doc, "References", style="Heading 1")
    for ref in data["references"]:
        add_docx_paragraph(doc, ieee_reference(ref), style="Normal")

    path = PUBLIC_DIR / "recircuit-research-whitepaper-v2.0.docx"
    doc.save(path)
    strip_docx_metadata(path)


def sync_legacy_pdf() -> None:
    source = PUBLIC_DIR / "recircuit-research-whitepaper-v2.0.pdf"
    legacy = ROOT / "public" / "recircuit-ai-circular-economy-whitepaper.pdf"
    shutil.copyfile(source, legacy)


def main() -> None:
    ensure_dirs()
    data = load_data()
    write_companion_docs(data)
    build_svgs()
    build_pdf(data)
    build_docx(data)
    sync_legacy_pdf()
    print("research artifacts generated")


if __name__ == "__main__":
    main()
